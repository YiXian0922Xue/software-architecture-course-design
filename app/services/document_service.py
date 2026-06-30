import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class DocumentService:
    text_extensions = {".txt", ".md", ".csv", ".json", ".py", ".log"}

    def extract(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in self.text_extensions:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".docx":
            return self._docx(path)
        if suffix == ".pdf":
            return "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
        if suffix in {".tex", ".latex"}:
            return self._latex(path)
        raise ValueError(f"暂不支持解析 {suffix or '无扩展名'} 文件")

    @staticmethod
    def _docx(path: Path) -> str:
        doc = Document(path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            lines.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(lines)

    @staticmethod
    def _latex(path: Path) -> str:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        raw = re.sub(r"%.*", "", raw)
        raw = re.sub(r"\\(?:begin|end)\{[^}]+\}", "\n", raw)
        raw = re.sub(r"\\(?:section|subsection|subsubsection|chapter)\*?\{([^}]*)\}", r"\n\1\n", raw)
        raw = re.sub(r"\\[a-zA-Z@]+(?:\[[^]]*\])?\{([^{}]*)\}", r"\1", raw)
        raw = re.sub(r"\\[a-zA-Z@]+|[{}]", " ", raw)
        return re.sub(r"\n{3,}", "\n\n", raw).strip()

